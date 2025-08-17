import os
import logging
import tempfile
import zipfile
from telegram import InlineKeyboardButton, InlineKeyboardMarkup, Update
from telegram.ext import (
    Application, CommandHandler, MessageHandler, CallbackQueryHandler,
    ContextTypes, filters
)
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import filetype
from dotenv import load_dotenv

# Load token from .env
load_dotenv()
BOT_TOKEN = os.getenv("BOT_TOKEN")

# Enable logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

CONVERT_OPTIONS = [
    ("📄 PDF", "pdf"),
    ("📃 DOCX", "docx"),
    ("🖼️ JPG", "jpg"),
    ("🖼️ PNG", "png"),
    ("📁 TXT", "txt"),
    ("🖼️ WEBP", "webp"),
    ("🖼️ BMP", "bmp"),
    ("📁 ZIP", "zip")
]

WELCOME_MESSAGE = (
    "👋 Hello! I'm **LokeshBotConverter**.\n"
    "Send me any file (PDF, Image, DOCX, etc.) and I'll convert it to the format you choose.\n\n"
    "✨ *Created with ❤️ by Lokesh.R*"
)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(WELCOME_MESSAGE, parse_mode='Markdown')

async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    file = update.message.document or (update.message.photo[-1] if update.message.photo else None)
    if not file:
        await update.message.reply_text("⚠️ Unsupported file type. Please send a document or an image.")
        return

    file_id = file.file_id
    context.user_data['file_id'] = file_id
    context.user_data['file_name'] = getattr(update.message.document, 'file_name', 'photo.jpg')

    keyboard = [
        [InlineKeyboardButton(opt[0], callback_data=opt[asset:1]) for opt in CONVERT_OPTIONS[i:i+2]]
        for i in range(0, len(CONVERT_OPTIONS), 2)
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("✅ Choose a format to convert:", reply_markup=reply_markup)

async def convert_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    format_chosen = query.data
    file_id = context.user_data.get("file_id")

    if not file_id:
        await query.edit_message_text("⚠️ No file to convert. Please send a new file.")
        return

    await query.edit_message_text("⏳ Converting file...")

    telegram_file = await context.bot.get_file(file_id)
    original_path, converted_path = "", ""

    try:
        with tempfile.NamedTemporaryFile(delete=False) as temp:
            await telegram_file.download_to_drive(custom_path=temp.name)
            original_path = temp.name

        converted_path = await perform_conversion(original_path, format_chosen)

        if converted_path:
            with open(converted_path, "rb") as f:
                file_name_base = os.path.splitext(context.user_data.get('file_name', 'file'))[0]
                await query.message.reply_document(f, filename=f"{file_name_base}.{format_chosen}")
            await query.message.reply_text("✅ Conversion successful!")
            await query.message.reply_text("✨ *Created with ❤️ by Lokesh.R*", parse_mode='Markdown')
        else:
            await query.edit_message_text("❌ Conversion failed. The selected format is not supported for this file type.")
    except Exception as e:
        logger.error(f"Error during conversion: {e}")
        await query.edit_message_text("❌ An error occurred during conversion. Please try again or use a different file.")
    finally:
        if os.path.exists(original_path): os.remove(original_path)
        if converted_path and os.path.exists(converted_path): os.remove(converted_path)
        context.user_data.clear()

async def perform_conversion(input_path, output_format):
    output_path = f"{input_path}.{output_format}"
    kind = filetype.guess(input_path)
    mime = kind.mime if kind else "unknown"

    try:
        if output_format in ["jpg", "png", "bmp", "webp"]:
            if mime.startswith("image/"):
                image = Image.open(input_path).convert("RGB")
                image.save(output_path, format=output_format.upper())
            elif mime == "application/pdf":
                doc = fitz.open(input_path)
                page = doc.load_page(0)
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img.save(output_path)
                doc.close()
            else:
                return None

        elif output_format == "pdf":
            if mime.startswith("image/"):
                image = Image.open(input_path).convert("RGB")
                image.save(output_path, "PDF", resolution=100.0)
            else:
                return None

        elif output_format == "docx":
            doc = Document()
            if mime.startswith("text/"):
                with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                    for line in f:
                        doc.add_paragraph(line.strip())
            elif mime.startswith("image/"):
                doc.add_picture(input_path, width=Inches(6.0))
            elif mime == "application/pdf":
                pdf_doc = fitz.open(input_path)
                for page in pdf_doc:
                    text = page.get_text()
                    doc.add_paragraph(text)
                pdf_doc.close()
            else:
                return None
            doc.save(output_path)

        elif output_format == "txt":
            if mime == "application/pdf":
                pdf_doc = fitz.open(input_path)
                text = ""
                for page in pdf_doc:
                    text += page.get_text()
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(text)
                pdf_doc.close()
            elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(input_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(text)
            else:
                return None

        elif output_format == "zip":
            with zipfile.ZipFile(output_path, 'w') as zipf:
                zipf.write(input_path, arcname=os.path.basename(input_path))
        else:
            return None
        return output_path
    except Exception as e:
        logger.error(f"Conversion logic failed for {input_path} to {output_format}: {e}")
        return None

def main():
    app = Application.builder().token(BOT_TOKEN).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.Document.ALL | filters.PHOTO, handle_file))
    app.add_handler(CallbackQueryHandler(convert_file))
    print("✅ Bot running as LokeshBotConverter...")
    app.run_polling()

if __name__ == "__main__":
    main()
